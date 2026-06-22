import json
import logging
import os
import re

import anthropic
from docx import Document

from app.config import settings
from app.services.resume_parser import extract_resume_text

logger = logging.getLogger(__name__)
client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

TAILOR_SYSTEM = """You are a professional resume writer. You will receive a user's base resume and a job description.
Your task is to rewrite the resume tailored to the specific job posting.

Rules:
- Adjust the professional summary to match the job's requirements
- Rewrite bullet points to use keywords and language from the job description
- Reorder skills to prioritize those mentioned in the job posting
- Keep all facts truthful — do not fabricate experience, just reframe existing experience
- Maintain the same general structure (sections, chronological order)
- Output the result as JSON with this structure:

{
  "name": "Full Name",
  "contact": "email | phone | location",
  "summary": "Professional summary paragraph",
  "experience": [
    {
      "title": "Job Title",
      "company": "Company Name",
      "dates": "Start - End",
      "bullets": ["bullet 1", "bullet 2", "bullet 3"]
    }
  ],
  "education": [
    {
      "degree": "Degree",
      "school": "School Name",
      "dates": "Graduation year or date range"
    }
  ],
  "skills": ["skill1", "skill2", "skill3"],
  "certifications": ["cert1", "cert2"]
}

Return ONLY valid JSON, no markdown fences or explanation."""


async def tailor_resume(base_resume_path: str, job_description: str, job_title: str, company: str) -> dict:
    resume_text = extract_resume_text(base_resume_path)

    prompt = f"""Base Resume:
{resume_text}

---

Target Job Posting:
Title: {job_title}
Company: {company}

Description:
{job_description}

---

Rewrite the resume to be tailored for this specific job posting."""

    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4000,
        system=TAILOR_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )

    from app.services.interview import _parse_json_response
    return _parse_json_response(message.content[0].text)


def create_docx(resume_data: dict, output_path: str) -> str:
    doc = Document()

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(resume_data.get("name", ""))
    name_run.bold = True
    name_run.font.size = shared_pt(16)
    name_para.alignment = 1  # Center

    # Contact
    if resume_data.get("contact"):
        contact_para = doc.add_paragraph(resume_data["contact"])
        contact_para.alignment = 1

    # Summary
    if resume_data.get("summary"):
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(resume_data["summary"])

    # Experience
    if resume_data.get("experience"):
        doc.add_heading("Experience", level=2)
        for exp in resume_data["experience"]:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.bold = True
            if exp.get("dates"):
                title_para.add_run(f"  |  {exp['dates']}")

            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # Education
    if resume_data.get("education"):
        doc.add_heading("Education", level=2)
        for edu in resume_data["education"]:
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(f"{edu.get('degree', '')} — {edu.get('school', '')}")
            edu_run.bold = True
            if edu.get("dates"):
                edu_para.add_run(f"  |  {edu['dates']}")

    # Skills
    if resume_data.get("skills"):
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(resume_data["skills"]))

    # Certifications
    if resume_data.get("certifications"):
        doc.add_heading("Certifications", level=2)
        for cert in resume_data["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def shared_pt(pt):
    from docx.shared import Pt
    return Pt(pt)


def sanitize_filename(s: str) -> str:
    return re.sub(r'[^\w\s-]', '', s).strip().replace(' ', '_')[:50]
